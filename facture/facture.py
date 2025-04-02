import datetime

def facture_txt(liste_produits, total_general):
    with open("facture.txt", "w", encoding="utf-8") as f:
        f.write("KING'S STORE :\nProduits achetés :\n")
        for produit in liste_produits:
            f.write(f"{produit['nom']}, \n - Quantité : {produit['quantite']}, Prix unitaire : {produit['prix_unitaire']}€, Total : {produit['total']}€\n")
        f.write(f"Total général : {total_general} €.\n")
        f.write(f"Date : {datetime.datetime.now().strftime('%d/%m/%Y - %H:%M')}\n")
        f.write("Merci de votre achat !\n")
    print("Facture générée sous format TXT.")
    
    
    
# Format de la facture Word
def facture_word(liste_produits, total_general):
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    document = Document()
    document.add_heading("KING'S STORE", 0)
    document.add_heading("Produits achetés :", level=3)
    
    style_produit = "Normal"
    
    for produit in liste_produits:
        document.add_paragraph(f"{produit['nom']}, \n - Quantité : {produit['quantite']}, Prix unitaire : {produit['prix_unitaire']}€, Total : {produit['total']}€", style=style_produit)
    document.add_paragraph("\n")
    document.add_paragraph(f"Total général : {total_general} €.")
    document.add_paragraph(f"Date : {datetime.datetime.now().strftime('%d/%m/%Y - %H:%M')}")
    document.save("facture.docx")
    print("Facture enregistrée sous format Word.")
    
    
# Format de la facture PDF
from fpdf import FPDF

def facture_pdf(liste_produits, total_general):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # 🏪 Titre
    pdf.cell(200, 10, txt="KING'S STORE", ln=1, align="C")
    pdf.cell(200, 10, txt="Produits achetés", ln=1, align="C")

    # 📦 Liste des produits
    for produit in liste_produits:
        texte = f"{produit['nom']}, \n  - Quantité : {produit['quantite']}, Prix unitaire : {produit['prix_unitaire']}€, Total : {produit['total']}€"
        
        # ✅ Remplacement des caractères spéciaux
        texte = texte.encode("latin-1", "replace").decode("latin-1")  
        
        pdf.set_font("Arial", "", 12)  # Assurer la police avant chaque texte
        pdf.multi_cell(0, 10, txt=texte)  # multi_cell gère les longues lignes
    
    # 💰 Total général
    total_texte = f"Total général : {total_general} €."
    total_texte = total_texte.encode("latin-1", "replace").decode("latin-1")  # Sécuriser l'encodage
    
   
    
    pdf.set_font("Arial", "", 12)
    pdf.cell(200, 10, txt=total_texte, ln=1, align="C")
    
    pdf.image("img1.jpg", x=10, y=8, w=33)  # Ajouter le logo
    
    date = datetime.datetime.now().strftime("%d/%m/%Y - %H:%M")
    date_texte = f"Date : {date}"
    date_texte = date_texte.encode("latin-1", "replace").decode("latin-1")  # Sécuriser l'encodage
    pdf.cell(200, 10, txt=date_texte, ln=1, align="C")

    # 📄 Sauvegarde du fichier PDF
    pdf.output("facture.pdf", "F")
    print("✅ Facture enregistrée sous format PDF.")


# Format de la facture Excel
def facture_excel(liste_produits, total_general):
    import pandas as pd

    df = pd.DataFrame(liste_produits)

    # Crée une nouvelle ligne avec le total général.  Adaptez les clés selon la structure de votre liste_produits.
    total_row = {'Nom du produit': 'TOTAL GENERAL', 'Prix': '', 'Quantité': '', 'Total': total_general} # Adaptez les noms de colonnes si nécessaire

    # Convertit la nouvelle ligne en Series et l'ajoute au DataFrame
    df = pd.concat([df, pd.DataFrame([total_row])], ignore_index=True)


    df.to_excel("facture.xlsx", index=False)
    print("Facture enregistrée sous format Excel.")
